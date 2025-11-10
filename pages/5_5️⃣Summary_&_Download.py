import streamlit as st
from core.auth import require_password
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Summary & Download", page_icon="📦", layout="wide")
require_password()

st.title("📦 Step 5 — Summary & Download")

# ---- Pull state from earlier steps ----
case_md = st.session_state.get("case_markdown")
cover_image = st.session_state.get("cover_image")                  # bytes (optional)
diagram_images = st.session_state.get("diagram_images", {})        # {prompt: bytes}

if not case_md:
    st.warning("Please complete **Step 3 – Draft Case Study** and **Step 4 – Generate Visuals** first.")
    st.stop()

# ---- Parse case study into sections (simple markdown parser) ----
# We expect your generated draft uses these H2 headings:
# ## Executive Summary, ## Problem / Need, ## Implementation Approach, ## Benefits & Impact, ## Key Learning Points, ## Point of Contact, ## Suggested Visuals / Diagrams
SECTION_ORDER = [
    "Title",
    "Executive Summary",
    "Problem / Need",
    "Implementation Approach",
    "Benefits & Impact",
    "Key Learning Points",
    "Point of Contact",
]

def parse_sections(md: str):
    sections = {"Title": []}
    current = "Title"
    for raw_line in md.splitlines():
        line = raw_line.strip()

        # H1 -> Title content
        if line.startswith("# "):
            sections["Title"] = [line[2:].strip()]
            current = "Title"
            continue

        # H2 -> start new section
        if line.startswith("## "):
            heading = line[3:].strip()
            # normalize to known section names if possible
            normalized = None
            for std in SECTION_ORDER[1:]:
                if std.lower().split()[0] in heading.lower():
                    normalized = std
                    break
            heading_key = normalized or heading
            sections.setdefault(heading_key, [])
            current = heading_key
            continue

        sections.setdefault(current, []).append(raw_line)  # preserve raw text for styling
    return sections

sections = parse_sections(case_md)

# Ensure all expected sections exist (even if empty)
for sec in SECTION_ORDER:
    sections.setdefault(sec, [])

# ---- Auto-suggest placements for diagram images (heuristics) ----
def suggest_section_for_prompt(prompt: str) -> str:
    p = prompt.lower()
    if any(k in p for k in ["exec", "overview", "summary"]):
        return "Executive Summary"
    if any(k in p for k in ["problem", "need", "challenge", "why"]):
        return "Problem / Need"
    if any(k in p for k in ["implement", "workflow", "process", "timeline", "approach"]):
        return "Implementation Approach"
    if any(k in p for k in ["benefit", "impact", "kpi", "metric", "outcome"]):
        return "Benefits & Impact"
    if any(k in p for k in ["learning", "lesson", "retrospective"]):
        return "Key Learning Points"
    if any(k in p for k in ["contact", "poc"]):
        return "Point of Contact"
    if any(k in p for k in ["diagram", "visual"]):
        return "Suggested Visuals / Diagrams"
    # default: Implementation (most diagrams are flows)
    return "Implementation Approach"

# Persist a mapping {image_key -> section_name}
# image_key for cover: "__COVER__"
if "image_placement" not in st.session_state:
    mapping = {}
    if cover_image:
        mapping["__COVER__"] = "Title"
    for prompt in diagram_images.keys():
        mapping[prompt] = suggest_section_for_prompt(prompt)
    st.session_state["image_placement"] = mapping

# ---- UI: preview + placement controls ----
st.header("Preview")
if cover_image:
    st.image(cover_image, caption="Cover Image (currently assigned to: Title) ", use_column_width=True)

st.markdown(case_md)

st.divider()
st.header("Image Placement")

with st.expander("Cover Image Placement", expanded=True if cover_image else False):
    if cover_image:
        st.session_state["image_placement"]["__COVER__"] = st.selectbox(
            "Place cover image under section:",
            options=SECTION_ORDER,
            index=SECTION_ORDER.index(st.session_state["image_placement"]["__COVER__"]),
            key="cover_select",
        )
    else:
        st.caption("No cover image generated in Step 4.")

if diagram_images:
    st.subheader("Diagram Placement")
    for prompt, img_bytes in diagram_images.items():
        cols = st.columns([3, 2])
        with cols[0]:
            st.image(img_bytes, caption=prompt, use_column_width=True)
        with cols[1]:
            default_sec = st.session_state["image_placement"].get(prompt, suggest_section_for_prompt(prompt))
            st.session_state["image_placement"][prompt] = st.selectbox(
                f"Place diagram for:\n“{prompt[:60]}{'…' if len(prompt)>60 else ''}”",
                options=SECTION_ORDER,
                index=SECTION_ORDER.index(default_sec) if default_sec in SECTION_ORDER else 3,
                key=f"sel_{hash(prompt)}",
            )
else:
    st.caption("No diagrams generated in Step 4.")

# ---- Build DOCX according to placements ----
def build_docx_from_sections(sections: dict, placement: dict, cover_img: bytes | None, diag_imgs: dict):
    doc = Document()

    # Helper: insert an image with sane width
    def add_img(img_bytes: bytes):
        doc.add_picture(BytesIO(img_bytes), width=Inches(5.5))
        doc.add_paragraph()

    # Title (H1)
    title_text = "Case Study"
    if sections.get("Title"):
        # First non-empty line used as title if H1 was present
        t = sections["Title"][0].strip("# ").strip()
        if t:
            title_text = t
    doc.add_heading(title_text, level=1)

    # If cover image is assigned to Title, add it now
    if cover_img and placement.get("__COVER__") == "Title":
        add_img(cover_img)

    # Iterate in our preferred order, then any extra sections the model may have added
    rendered = set(["Title"])
    ordered_keys = [k for k in SECTION_ORDER if k != "Title"] + [k for k in sections.keys() if k not in SECTION_ORDER]
    for key in ordered_keys:
        if key in rendered:
            continue
        rendered.add(key)

        # Heading
        doc.add_heading(key, level=2)

        # Body
        body_lines = sections.get(key, [])
        # naive markdown-ish rendering
        for raw in body_lines:
            line = raw.rstrip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                doc.add_paragraph(line.strip("*"), style="Intense Quote")
            else:
                doc.add_paragraph(line)

        # Any images mapped to this section?
        # Cover image (if not Title)
        if cover_img and placement.get("__COVER__") == key and key != "Title":
            add_img(cover_img)

        # Diagrams assigned here
        for prompt, img in (diag_imgs or {}).items():
            if placement.get(prompt) == key:
                add_img(img)

    # If any images still unmapped (unlikely), append them at the end
    for prompt, img in (diag_imgs or {}).items():
        if placement.get(prompt) not in sections:
            doc.add_page_break()
            doc.add_heading("Additional Visual", level=2)
            doc.add_paragraph(prompt)
            add_img(img)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.divider()
if st.button("📥 Generate & Download DOCX"):
    doc_buf = build_docx_from_sections(
        sections=sections,
        placement=st.session_state["image_placement"],
        cover_img=cover_image,
        diag_imgs=diagram_images,
    )
    st.success("✅ DOCX generated successfully!")
    st.download_button(
        label="⬇️ Download Case Study (.docx)",
        data=doc_buf,
        file_name="LEAPscribe_Case_Study.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
